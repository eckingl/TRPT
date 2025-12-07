"""属性图报告生成模块

整合数据处理、图表生成、AI 分析和 Word 输出，生成完整的属性图分析报告
"""

from dataclasses import dataclass
from datetime import datetime
from io import BytesIO
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Mm, Pt

from app.core.chart import (
    make_grade_bar_chart,
    make_grade_pie_chart,
    make_town_comparison_chart,
    make_town_grade_stack_chart,
    set_theme,
)
from app.topics.attribute_map.config import SOIL_ATTR_CONFIG, get_grade_order
from app.topics.attribute_map.stats import AttributeStats


@dataclass
class ReportConfig:
    """报告配置"""

    region_name: str = "XX县"
    survey_year: int = 2024
    theme: str = "professional"
    image_width_mm: int = 150
    include_pie_chart: bool = True
    include_bar_chart: bool = True
    include_town_chart: bool = True
    include_stack_chart: bool = True
    # AI 相关配置
    use_ai: bool = False
    ai_provider: str | None = None  # qwen / deepseek，None 使用默认


def generate_attribute_report(
    stats: AttributeStats,
    config: ReportConfig | None = None,
    output_path: Path | None = None,
) -> bytes:
    """生成单个属性的分析报告

    报告结构：
    一、总体情况（样点分析 + 制图分析 + 溯源分析 + 总体情况表）
    二、土地利用类型分析
    三、土壤类型分析
    四、乡镇分析

    Args:
        stats: 属性统计数据
        config: 报告配置
        output_path: 输出路径

    Returns:
        bytes: Word 文档字节数据
    """
    if config is None:
        config = ReportConfig()

    set_theme(config.theme)

    doc = Document()

    # 文档标题
    title = f"{config.region_name}{stats.attr_name}分析报告"
    title_para = doc.add_heading(title, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 一、总体情况
    _add_overall_situation(doc, stats, config)

    # 二、土地利用类型分析
    _add_land_use_analysis(doc, stats, config)

    # 三、土壤类型分析
    _add_soil_type_analysis(doc, stats, config)

    # 四、乡镇分析
    _add_town_analysis_section(doc, stats, config)

    # 生成时间
    doc.add_paragraph()
    time_para = doc.add_paragraph()
    time_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = time_para.add_run(f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    run.font.size = Pt(9)
    run.font.italic = True

    # 保存
    return _save_document(doc, output_path)


def generate_multi_attribute_report(
    stats_list: list[AttributeStats],
    config: ReportConfig | None = None,
    output_path: Path | None = None,
) -> bytes:
    """生成多个属性的综合分析报告

    Args:
        stats_list: 属性统计数据列表
        config: 报告配置
        output_path: 输出路径

    Returns:
        bytes: Word 文档字节数据
    """
    if config is None:
        config = ReportConfig()

    set_theme(config.theme)

    doc = Document()

    # 文档标题
    title = f"{config.region_name}土壤属性综合分析报告"
    title_para = doc.add_heading(title, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 报告概述
    doc.add_heading("一、报告概述", level=1)
    intro_para = doc.add_paragraph()
    intro_para.add_run(
        f"本报告基于{config.region_name}{config.survey_year}年土壤普查数据，"
        f"对{len(stats_list)}个土壤属性进行统计分析。"
    )

    # 属性列表
    attr_list = "、".join([s.attr_name for s in stats_list])
    doc.add_paragraph(f"分析属性包括：{attr_list}。")

    # AI 综合摘要（如果启用）
    if config.use_ai:
        _add_ai_comprehensive_summary(doc, stats_list, config)

    # 各属性分析
    for i, stats in enumerate(stats_list, 1):
        doc.add_page_break()
        _add_attribute_section(doc, stats, config, section_num=i + 1)

    # 生成时间
    doc.add_paragraph()
    time_para = doc.add_paragraph()
    time_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = time_para.add_run(f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    run.font.size = Pt(9)
    run.font.italic = True

    return _save_document(doc, output_path)


def generate_single_attribute_reports(
    stats_list: list[AttributeStats],
    config: ReportConfig | None = None,
    output_dir: Path | None = None,
) -> list[tuple[str, bytes]]:
    """为每个属性单独生成报告

    Args:
        stats_list: 属性统计数据列表
        config: 报告配置
        output_dir: 输出目录

    Returns:
        list[tuple[str, bytes]]: 列表，每项为 (文件名, 文档字节数据)
    """
    if config is None:
        config = ReportConfig()

    results = []
    for stats in stats_list:
        filename = f"{config.region_name}_{stats.attr_name}_分析报告.docx"
        output_path = output_dir / filename if output_dir else None

        data = generate_attribute_report(stats, config, output_path)
        results.append((filename, data))

    return results


def _add_basic_info(doc: Document, stats: AttributeStats, config: ReportConfig) -> None:
    """添加基本信息"""
    doc.add_heading("一、基本信息", level=1)

    table = doc.add_table(rows=4, cols=2)
    table.style = "Table Grid"

    rows_data = [
        ("地区", config.region_name),
        ("调查年份", str(config.survey_year)),
        ("属性名称", stats.attr_name),
        ("单位", stats.unit),
    ]

    for i, (label, value) in enumerate(rows_data):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = value
        # 标签加粗
        for para in table.rows[i].cells[0].paragraphs:
            for run in para.runs:
                run.font.bold = True

    doc.add_paragraph()


def _add_statistics_summary(doc: Document, stats: AttributeStats) -> None:
    """添加统计概要"""
    doc.add_heading("二、统计概要", level=1)

    # 样点统计
    doc.add_heading("2.1 样点数据统计", level=2)

    table = doc.add_table(rows=5, cols=2)
    table.style = "Table Grid"

    sample_data = [
        ("样点总数", f"{stats.sample_total}"),
        ("平均值", f"{stats.sample_mean:.2f} {stats.unit}"),
        ("中位值", f"{stats.sample_median:.2f} {stats.unit}"),
        ("最小值", f"{stats.sample_min:.2f} {stats.unit}"),
        ("最大值", f"{stats.sample_max:.2f} {stats.unit}"),
    ]

    for i, (label, value) in enumerate(sample_data):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = value

    doc.add_paragraph()

    # 制图数据统计
    if stats.area_total > 0:
        doc.add_heading("2.2 制图数据统计", level=2)

        table2 = doc.add_table(rows=5, cols=2)
        table2.style = "Table Grid"

        area_data = [
            ("总面积", f"{stats.area_total:.2f} 亩"),
            ("平均值", f"{stats.area_mean:.2f} {stats.unit}"),
            ("中位值", f"{stats.area_median:.2f} {stats.unit}"),
            ("最小值", f"{stats.area_min:.2f} {stats.unit}"),
            ("最大值", f"{stats.area_max:.2f} {stats.unit}"),
        ]

        for i, (label, value) in enumerate(area_data):
            table2.rows[i].cells[0].text = label
            table2.rows[i].cells[1].text = value

        doc.add_paragraph()


def _add_ai_analysis(
    doc: Document, stats: AttributeStats, config: ReportConfig
) -> None:
    """添加 AI 生成的分析内容"""
    from app.core.ai import generate_analysis

    doc.add_heading("2.3 专家分析", level=2)

    # 计算等级分布百分比
    total = sum(stats.grade_area_sums.values())
    grade_pct = {}
    if total > 0:
        grade_pct = {k: v / total * 100 for k, v in stats.grade_area_sums.items()}

    try:
        analysis_text = generate_analysis(
            attr_name=stats.attr_name,
            unit=stats.unit,
            sample_total=stats.sample_total,
            sample_mean=stats.sample_mean,
            sample_median=stats.sample_median,
            sample_min=stats.sample_min,
            sample_max=stats.sample_max,
            grade_distribution=grade_pct,
            region_name=config.region_name,
            provider=config.ai_provider,
        )
        doc.add_paragraph(analysis_text)
    except Exception as e:
        doc.add_paragraph(f"[AI 分析生成失败: {e}]")

    doc.add_paragraph()


def _add_grade_distribution(
    doc: Document, stats: AttributeStats, config: ReportConfig
) -> None:
    """添加等级分布分析"""
    doc.add_heading("三、等级分布分析", level=1)

    # 等级分布表格
    doc.add_heading("3.1 等级统计表", level=2)

    # 构建表格数据
    grade_config = SOIL_ATTR_CONFIG.get(stats.attr_key, {})
    levels = grade_config.get("levels", [])

    headers = ["等级", "值域", "样点数", "样点占比", "面积(亩)", "面积占比"]
    rows = []

    sample_total = sum(stats.grade_sample_counts.values())
    area_total = sum(stats.grade_area_sums.values())

    for _threshold, grade_name, desc in levels:
        sample_count = stats.grade_sample_counts.get(grade_name, 0)
        area_sum = stats.grade_area_sums.get(grade_name, 0)

        sample_pct = sample_count / sample_total * 100 if sample_total > 0 else 0
        area_pct = area_sum / area_total * 100 if area_total > 0 else 0

        rows.append(
            [
                grade_name,
                desc,
                str(sample_count),
                f"{sample_pct:.1f}%",
                f"{area_sum:.1f}",
                f"{area_pct:.1f}%",
            ]
        )

    _insert_table(doc, headers, rows)
    doc.add_paragraph()

    # 饼图
    if config.include_pie_chart and area_total > 0:
        doc.add_heading("3.2 等级分布饼图", level=2)
        pie_data = make_grade_pie_chart(
            stats.grade_area_sums,
            f"{stats.attr_name}等级分布",
        )
        _insert_image(doc, pie_data, config.image_width_mm, "图：等级面积分布")

    # 柱状图
    if config.include_bar_chart and area_total > 0:
        doc.add_heading("3.3 等级分布柱状图", level=2)
        bar_data = make_grade_bar_chart(
            stats.grade_area_sums,
            f"{stats.attr_name}分级面积",
            ylabel="面积(亩)",
        )
        _insert_image(doc, bar_data, config.image_width_mm, "图：各等级面积对比")


def _add_town_analysis(
    doc: Document, stats: AttributeStats, config: ReportConfig
) -> None:
    """添加乡镇分析"""
    doc.add_heading("四、乡镇分析", level=1)

    # 乡镇对比图
    if config.include_town_chart:
        doc.add_heading("4.1 乡镇均值对比", level=2)
        town_chart = make_town_comparison_chart(stats.town_stats, stats.attr_name)
        _insert_image(doc, town_chart, config.image_width_mm, "图：各乡镇均值对比")

    # 堆叠图
    if config.include_stack_chart:
        grade_order = get_grade_order(stats.attr_key)

        doc.add_heading("4.2 乡镇等级分布", level=2)
        stack_chart = make_town_grade_stack_chart(
            stats.town_stats,
            grade_order,
            stats.attr_name,
        )
        _insert_image(doc, stack_chart, config.image_width_mm, "图：各乡镇等级占比分布")


def _add_ai_conclusion_and_suggestion(
    doc: Document, stats: AttributeStats, config: ReportConfig
) -> None:
    """添加 AI 生成的结论和建议"""
    from app.core.ai import generate_conclusion, generate_suggestion

    doc.add_heading("五、结论与建议", level=1)

    # 计算等级分布百分比
    total = sum(stats.grade_area_sums.values())
    grade_pct = {}
    if total > 0:
        grade_pct = {k: v / total * 100 for k, v in stats.grade_area_sums.items()}

    # 结论
    doc.add_heading("5.1 分析结论", level=2)
    try:
        conclusion = generate_conclusion(
            attr_name=stats.attr_name,
            sample_mean=stats.sample_mean,
            unit=stats.unit,
            grade_distribution=grade_pct,
            region_name=config.region_name,
            provider=config.ai_provider,
        )
        doc.add_paragraph(conclusion)
    except Exception as e:
        doc.add_paragraph(f"[结论生成失败: {e}]")

    # 建议
    doc.add_heading("5.2 改良建议", level=2)
    try:
        suggestion = generate_suggestion(
            attr_name=stats.attr_name,
            sample_mean=stats.sample_mean,
            unit=stats.unit,
            grade_distribution=grade_pct,
            region_name=config.region_name,
            provider=config.ai_provider,
        )
        doc.add_paragraph(suggestion)
    except Exception as e:
        doc.add_paragraph(f"[建议生成失败: {e}]")

    doc.add_paragraph()


def _add_ai_comprehensive_summary(
    doc: Document, stats_list: list[AttributeStats], config: ReportConfig
) -> None:
    """添加 AI 生成的综合摘要"""
    from app.core.ai import generate_comprehensive_summary

    doc.add_heading("1.1 综合评价", level=2)

    # 构建各属性摘要
    attr_summaries = []
    for stats in stats_list:
        total = sum(stats.grade_area_sums.values())
        main_grade = "未知"
        if total > 0:
            main_grade = max(stats.grade_area_sums.items(), key=lambda x: x[1])[0]

        attr_summaries.append(
            {
                "name": stats.attr_name,
                "mean": stats.sample_mean,
                "unit": stats.unit,
                "main_grade": main_grade,
            }
        )

    try:
        summary = generate_comprehensive_summary(
            region_name=config.region_name,
            survey_year=config.survey_year,
            attr_summaries=attr_summaries,
            provider=config.ai_provider,
        )
        doc.add_paragraph(summary)
    except Exception as e:
        doc.add_paragraph(f"[综合摘要生成失败: {e}]")

    doc.add_paragraph()


def _add_attribute_section(
    doc: Document,
    stats: AttributeStats,
    config: ReportConfig,
    section_num: int,
) -> None:
    """添加单个属性分析章节"""
    # 章节标题
    doc.add_heading(f"{_to_chinese_num(section_num)}、{stats.attr_name}分析", level=1)

    # 统计概要
    doc.add_heading(f"{section_num}.1 统计概要", level=2)

    para = doc.add_paragraph()
    para.add_run(
        f"{stats.attr_name}（{stats.unit}）共有{stats.sample_total}个样点，"
        f"均值为{stats.sample_mean:.2f}，中位值为{stats.sample_median:.2f}，"
        f"范围为{stats.sample_min:.2f}~{stats.sample_max:.2f}。"
    )

    if stats.area_total > 0:
        para.add_run(f"制图总面积为{stats.area_total:.2f}亩。")

    # AI 分析（如果启用）
    if config.use_ai:
        from app.core.ai import generate_analysis

        total = sum(stats.grade_area_sums.values())
        grade_pct = (
            {k: v / total * 100 for k, v in stats.grade_area_sums.items()}
            if total > 0
            else {}
        )

        try:
            analysis = generate_analysis(
                attr_name=stats.attr_name,
                unit=stats.unit,
                sample_total=stats.sample_total,
                sample_mean=stats.sample_mean,
                sample_median=stats.sample_median,
                sample_min=stats.sample_min,
                sample_max=stats.sample_max,
                grade_distribution=grade_pct,
                region_name=config.region_name,
                provider=config.ai_provider,
            )
            doc.add_paragraph(analysis)
        except Exception:
            pass

    # 等级分布图
    area_total = sum(stats.grade_area_sums.values())
    if config.include_pie_chart and area_total > 0:
        doc.add_heading(f"{section_num}.2 等级分布", level=2)
        pie_data = make_grade_pie_chart(
            stats.grade_area_sums,
            f"{stats.attr_name}等级分布",
        )
        _insert_image(doc, pie_data, config.image_width_mm)

    # 乡镇对比
    if config.include_town_chart and not stats.town_stats.empty:
        doc.add_heading(f"{section_num}.3 乡镇对比", level=2)
        town_chart = make_town_comparison_chart(stats.town_stats, stats.attr_name)
        _insert_image(doc, town_chart, config.image_width_mm)


def _insert_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    """插入表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"

    # 表头
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.bold = True

    # 数据行
    for row_idx, row_data in enumerate(rows):
        for col_idx, value in enumerate(row_data):
            if col_idx < len(table.rows[row_idx + 1].cells):
                table.rows[row_idx + 1].cells[col_idx].text = value


def _insert_image(
    doc: Document,
    image_data: bytes,
    width_mm: int,
    caption: str | None = None,
) -> None:
    """插入图片"""
    if not image_data:
        return

    img_buf = BytesIO(image_data)
    pic_para = doc.add_paragraph()
    pic_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = pic_para.add_run()
    run.add_picture(img_buf, width=Mm(width_mm))

    if caption:
        caption_para = doc.add_paragraph()
        caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_run = caption_para.add_run(caption)
        cap_run.font.size = Pt(10)
        cap_run.font.italic = True

    doc.add_paragraph()


def _save_document(doc: Document, output_path: Path | None) -> bytes:
    """保存文档"""
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    data = buf.read()

    if output_path is not None:
        output_path.parent.mkdir(parents=True, exist_ok=True)
        output_path.write_bytes(data)

    return data


def _to_chinese_num(num: int) -> str:
    """数字转中文"""
    chinese_nums = ["零", "一", "二", "三", "四", "五", "六", "七", "八", "九", "十"]
    if num <= 10:
        return chinese_nums[num]
    elif num < 20:
        return f"十{chinese_nums[num - 10]}" if num > 10 else "十"
    else:
        return str(num)


def _add_overall_situation(
    doc: Document, stats: AttributeStats, config: ReportConfig
) -> None:
    """添加总体情况章节

    包含：样点分析 + 制图分析 + 溯源分析（连续段落）+ 总体情况表
    """
    doc.add_heading("一、总体情况", level=1)

    grade_config = SOIL_ATTR_CONFIG.get(stats.attr_key, {})
    levels = grade_config.get("levels", [])

    # 构建样点分析文本
    sample_parts = [
        f"{config.region_name}{stats.attr_name}样点统计分析：共采集有效样点{stats.sample_total}个，"
        f"样点{stats.attr_name}平均值为{stats.sample_mean:.2f}{stats.unit}，"
        f"中位值为{stats.sample_median:.2f}{stats.unit}，"
        f"数值范围为{stats.sample_min:.2f}~{stats.sample_max:.2f}{stats.unit}"
    ]

    # 等级分布描述
    if levels and stats.sample_total > 0:
        grade_desc_parts = []
        for _threshold, grade_name, _desc in levels:
            count = stats.grade_sample_counts.get(grade_name, 0)
            if count > 0:
                pct = count / stats.sample_total * 100
                grade_desc_parts.append(f"{grade_name}{count}个（占{pct:.1f}%）")
        if grade_desc_parts:
            sample_parts.append(
                f"，按等级划分，样点分布为：{'，'.join(grade_desc_parts)}"
            )

    sample_parts.append("。")

    # 样点分析段落
    doc.add_paragraph("".join(sample_parts))

    # 构建制图分析文本
    if stats.area_total > 0:
        area_parts = [
            f"{config.region_name}{stats.attr_name}制图统计分析：制图总面积为{stats.area_total:.2f}亩，"
            f"加权平均值为{stats.area_mean:.2f}{stats.unit}，"
            f"中位值为{stats.area_median:.2f}{stats.unit}，"
            f"数值范围为{stats.area_min:.2f}~{stats.area_max:.2f}{stats.unit}"
        ]

        # 等级面积分布
        total_area = sum(stats.grade_area_sums.values())
        if total_area > 0:
            area_grade_parts = []
            for _threshold, grade_name, _desc in levels:
                area_sum = stats.grade_area_sums.get(grade_name, 0)
                if area_sum > 0:
                    pct = area_sum / total_area * 100
                    area_grade_parts.append(
                        f"{grade_name}{area_sum:.1f}亩（占{pct:.1f}%）"
                    )
            if area_grade_parts:
                area_parts.append(
                    f"，按等级划分，面积分布为：{'，'.join(area_grade_parts)}"
                )

        area_parts.append("。")

        # 制图分析段落
        doc.add_paragraph("".join(area_parts))
    else:
        doc.add_paragraph("暂无制图数据。")

    # 溯源分析（AI生成，连续段落）
    if config.use_ai:
        _add_ai_traceability_analysis(doc, stats, config)

    doc.add_paragraph()

    # 总体情况统计表
    doc.add_heading("1.1 总体情况统计表", level=2)
    _add_overall_table(doc, stats)

    # 等级分布图表
    area_total = sum(stats.grade_area_sums.values())
    if config.include_pie_chart and area_total > 0:
        pie_data = make_grade_pie_chart(
            stats.grade_area_sums,
            f"{stats.attr_name}等级分布",
        )
        _insert_image(
            doc, pie_data, config.image_width_mm, f"图：{stats.attr_name}等级面积分布"
        )

    if config.include_bar_chart and area_total > 0:
        bar_data = make_grade_bar_chart(
            stats.grade_area_sums,
            f"{stats.attr_name}分级面积",
            ylabel="面积(亩)",
        )
        _insert_image(
            doc, bar_data, config.image_width_mm, f"图：{stats.attr_name}各等级面积对比"
        )


def _add_ai_traceability_analysis(
    doc: Document, stats: AttributeStats, config: ReportConfig
) -> None:
    """添加AI溯源分析"""
    from app.core.ai import generate_traceability_analysis

    # 计算等级分布百分比
    total = sum(stats.grade_area_sums.values())
    grade_pct = {}
    if total > 0:
        grade_pct = {k: v / total * 100 for k, v in stats.grade_area_sums.items()}

    try:
        analysis = generate_traceability_analysis(
            attr_name=stats.attr_name,
            unit=stats.unit,
            sample_mean=stats.sample_mean,
            area_mean=stats.area_mean,
            grade_distribution=grade_pct,
            region_name=config.region_name,
            provider=config.ai_provider,
        )
        doc.add_paragraph(analysis)
    except Exception as e:
        doc.add_paragraph(f"[溯源分析生成失败: {e}]")


def _add_overall_table(doc: Document, stats: AttributeStats) -> None:
    """添加总体情况统计表"""
    grade_config = SOIL_ATTR_CONFIG.get(stats.attr_key, {})
    levels = grade_config.get("levels", [])

    # 构建表格
    headers = ["等级", "值域", "样点数", "样点占比", "面积(亩)", "面积占比"]
    rows = []

    sample_total = stats.sample_total
    area_total = sum(stats.grade_area_sums.values())

    for _threshold, grade_name, desc in levels:
        sample_count = stats.grade_sample_counts.get(grade_name, 0)
        area_sum = stats.grade_area_sums.get(grade_name, 0)

        sample_pct = sample_count / sample_total * 100 if sample_total > 0 else 0
        area_pct = area_sum / area_total * 100 if area_total > 0 else 0

        rows.append(
            [
                grade_name,
                desc,
                str(sample_count),
                f"{sample_pct:.1f}%",
                f"{area_sum:.1f}",
                f"{area_pct:.1f}%",
            ]
        )

    _insert_table(doc, headers, rows)

    # 添加合计行说明
    summary = doc.add_paragraph()
    summary.add_run(
        f"合计：样点{sample_total}个，面积{area_total:.1f}亩；"
        f"均值：样点{stats.sample_mean:.2f}{stats.unit}，制图{stats.area_mean:.2f}{stats.unit}"
    )


def _add_land_use_analysis(
    doc: Document, stats: AttributeStats, config: ReportConfig
) -> None:
    """添加土地利用类型分析章节"""
    doc.add_heading("二、土地利用类型分析", level=1)

    df_sample = stats.df_sample_clean
    df_area = stats.df_area_clean
    attr_key = stats.attr_key

    # 检查是否有土地利用分类数据
    has_land_use_data = "一级" in df_sample.columns or "一级" in df_area.columns

    if not has_land_use_data:
        doc.add_paragraph("数据中未包含土地利用类型信息，无法进行分类分析。")
        return

    # 构建样点分析文本
    sample_text_parts = []
    if "一级" in df_sample.columns:
        land_use_sample = df_sample.groupby("一级", observed=True).agg(
            {attr_key: ["count", "mean", "median", "min", "max"]}
        )
        land_use_sample.columns = ["数量", "均值", "中位值", "最小值", "最大值"]

        sample_desc_parts = []
        for land_type, row in land_use_sample.iterrows():
            sample_desc_parts.append(
                f"{land_type}{int(row['数量'])}个（均值{row['均值']:.2f}{stats.unit}）"
            )
        if sample_desc_parts:
            sample_text_parts.append(
                f"按土地利用类型，样点分布为：{'，'.join(sample_desc_parts)}"
            )

    # 构建制图分析文本
    area_text_parts = []
    if "一级" in df_area.columns and "面积" in df_area.columns:
        land_use_area = df_area.groupby("一级", observed=True).agg(
            {"面积": "sum", attr_key: "mean"}
        )

        area_desc_parts = []
        for land_type, row in land_use_area.iterrows():
            area_desc_parts.append(
                f"{land_type}{row['面积']:.1f}亩（均值{row[attr_key]:.2f}{stats.unit}）"
            )
        if area_desc_parts:
            area_text_parts.append(
                f"按土地利用类型，面积分布为：{'，'.join(area_desc_parts)}"
            )

    # 合并为连续段落
    if sample_text_parts or area_text_parts:
        combined_text = "。".join(sample_text_parts + area_text_parts) + "。"
        doc.add_paragraph(combined_text)
    else:
        doc.add_paragraph("数据中无土地利用类型统计信息。")

    # AI 分析（如果启用，直接追加段落）
    if config.use_ai:
        _add_ai_land_use_analysis(doc, stats, config)

    doc.add_paragraph()

    # 土地利用类型统计表
    doc.add_heading("2.1 土地利用类型统计表", level=2)
    _add_land_use_table(doc, stats)


def _add_ai_land_use_analysis(
    doc: Document, stats: AttributeStats, config: ReportConfig
) -> None:
    """添加AI土地利用类型分析"""
    from app.core.ai import generate_land_use_analysis

    df_area = stats.df_area_clean
    attr_key = stats.attr_key

    # 构建土地利用数据摘要
    land_use_data = {}
    if "一级" in df_area.columns and "面积" in df_area.columns:
        grouped = df_area.groupby("一级", observed=True).agg(
            {"面积": "sum", attr_key: "mean"}
        )
        for land_type, row in grouped.iterrows():
            land_use_data[land_type] = {"area": row["面积"], "mean": row[attr_key]}

    try:
        analysis = generate_land_use_analysis(
            attr_name=stats.attr_name,
            unit=stats.unit,
            land_use_data=land_use_data,
            region_name=config.region_name,
            provider=config.ai_provider,
        )
        doc.add_paragraph(analysis)
    except Exception as e:
        doc.add_paragraph(f"[土地利用分析生成失败: {e}]")


def _add_land_use_table(doc: Document, stats: AttributeStats) -> None:
    """添加土地利用类型统计表"""
    df_sample = stats.df_sample_clean
    df_area = stats.df_area_clean
    attr_key = stats.attr_key

    headers = ["土地利用类型", "样点数", "样点均值", "面积(亩)", "面积均值"]
    rows = []

    land_types = set()
    if "一级" in df_sample.columns:
        land_types.update(df_sample["一级"].dropna().unique())
    if "一级" in df_area.columns:
        land_types.update(df_area["一级"].dropna().unique())

    for land_type in sorted(land_types):
        # 样点统计
        if "一级" in df_sample.columns:
            sample_data = df_sample[df_sample["一级"] == land_type]
            sample_count = len(sample_data)
            sample_mean = sample_data[attr_key].mean() if len(sample_data) > 0 else 0
        else:
            sample_count = 0
            sample_mean = 0

        # 面积统计
        if "一级" in df_area.columns and "面积" in df_area.columns:
            area_data = df_area[df_area["一级"] == land_type]
            area_sum = area_data["面积"].sum() if len(area_data) > 0 else 0
            area_mean = area_data[attr_key].mean() if len(area_data) > 0 else 0
        else:
            area_sum = 0
            area_mean = 0

        rows.append(
            [
                land_type,
                str(sample_count),
                f"{sample_mean:.2f}" if sample_mean else "-",
                f"{area_sum:.1f}",
                f"{area_mean:.2f}" if area_mean else "-",
            ]
        )

    _insert_table(doc, headers, rows)


def _add_soil_type_analysis(
    doc: Document, stats: AttributeStats, config: ReportConfig
) -> None:
    """添加土壤类型分析章节"""
    doc.add_heading("三、土壤类型分析", level=1)

    df = stats.soil_type_stats

    if df.empty:
        doc.add_paragraph("数据中未包含土壤类型信息（YL/TS），无法进行分类分析。")
        return

    # 构建样点分析文本
    sample_text_parts = []
    if "sample_count" in df.columns:
        soil_summary = df.groupby("YL", observed=True).agg(
            {"sample_count": "sum", "sample_mean": "mean"}
        )
        sample_desc_parts = []
        for yl, row in soil_summary.iterrows():
            if row["sample_count"] > 0:
                sample_desc_parts.append(
                    f"{yl}{int(row['sample_count'])}个（均值{row['sample_mean']:.2f}{stats.unit}）"
                )
        if sample_desc_parts:
            sample_text_parts.append(
                f"按土壤亚类，样点分布为：{'，'.join(sample_desc_parts)}"
            )

    # 构建制图分析文本
    area_text_parts = []
    if "area_sum" in df.columns:
        soil_area = df.groupby("YL", observed=True).agg(
            {"area_sum": "sum", "area_mean": "mean"}
        )
        area_desc_parts = []
        for yl, row in soil_area.iterrows():
            if row["area_sum"] > 0:
                area_desc_parts.append(
                    f"{yl}{row['area_sum']:.1f}亩（均值{row['area_mean']:.2f}{stats.unit}）"
                )
        if area_desc_parts:
            area_text_parts.append(
                f"按土壤亚类，面积分布为：{'，'.join(area_desc_parts)}"
            )

    # 合并为连续段落
    if sample_text_parts or area_text_parts:
        combined_text = "。".join(sample_text_parts + area_text_parts) + "。"
        doc.add_paragraph(combined_text)
    else:
        doc.add_paragraph("数据中无土壤类型统计信息。")

    # AI 分析（如果启用，直接追加段落）
    if config.use_ai:
        _add_ai_soil_type_analysis(doc, stats, config)

    doc.add_paragraph()

    # 土壤类型统计表
    doc.add_heading("3.1 土壤类型统计表", level=2)
    _add_soil_type_table(doc, stats)


def _add_ai_soil_type_analysis(
    doc: Document, stats: AttributeStats, config: ReportConfig
) -> None:
    """添加AI土壤类型分析"""
    from app.core.ai import generate_soil_type_analysis

    df = stats.soil_type_stats

    # 构建土壤类型数据摘要
    soil_data = {}
    if "area_sum" in df.columns:
        grouped = df.groupby("YL", observed=True).agg(
            {"area_sum": "sum", "area_mean": "mean"}
        )
        for yl, row in grouped.iterrows():
            soil_data[yl] = {"area": row["area_sum"], "mean": row["area_mean"]}

    try:
        analysis = generate_soil_type_analysis(
            attr_name=stats.attr_name,
            unit=stats.unit,
            soil_data=soil_data,
            region_name=config.region_name,
            provider=config.ai_provider,
        )
        doc.add_paragraph(analysis)
    except Exception as e:
        doc.add_paragraph(f"[土壤类型分析生成失败: {e}]")


def _add_soil_type_table(doc: Document, stats: AttributeStats) -> None:
    """添加土壤类型统计表"""
    df = stats.soil_type_stats

    if df.empty:
        doc.add_paragraph("无土壤类型统计数据。")
        return

    headers = ["亚类", "土属", "样点数", "样点均值", "面积(亩)", "面积均值"]
    rows = []

    for _, row in df.iterrows():
        yl = row.get("YL", "-")
        ts = row.get("TS", "-")
        sample_count = int(row.get("sample_count", 0)) if row.get("sample_count") else 0
        sample_mean = f"{row['sample_mean']:.2f}" if row.get("sample_mean") else "-"
        area_sum = f"{row['area_sum']:.1f}" if row.get("area_sum") else "-"
        area_mean = f"{row['area_mean']:.2f}" if row.get("area_mean") else "-"

        rows.append([yl, ts, str(sample_count), sample_mean, area_sum, area_mean])

    _insert_table(doc, headers, rows)


def _add_town_analysis_section(
    doc: Document, stats: AttributeStats, config: ReportConfig
) -> None:
    """添加乡镇分析章节"""
    doc.add_heading("四、乡镇分析", level=1)

    if stats.town_stats.empty:
        doc.add_paragraph("数据中未包含乡镇信息（行政区名称），无法进行分类分析。")
        return

    df = stats.town_stats

    # 构建样点分析文本
    sample_desc_parts = []
    for _, row in df.iterrows():
        town = row["乡镇"]
        count = int(row["样点数"])
        mean_val = row["均值"]
        if count > 0:
            mean_str = (
                f"均值{mean_val:.2f}{stats.unit}"
                if mean_val and mean_val == mean_val
                else ""
            )
            sample_desc_parts.append(
                f"{town}{count}个{'（' + mean_str + '）' if mean_str else ''}"
            )

    # 构建制图分析文本
    area_desc_parts = []
    for _, row in df.iterrows():
        town = row["乡镇"]
        area = row["面积"]
        if area > 0:
            area_desc_parts.append(f"{town}{area:.1f}亩")

    # 合并为连续段落
    text_parts = []
    if sample_desc_parts:
        sample_text = f"按乡镇，样点分布为：{'，'.join(sample_desc_parts[:10])}{'等' if len(sample_desc_parts) > 10 else ''}"
        text_parts.append(sample_text)
    if area_desc_parts:
        area_text = f"按乡镇，面积分布为：{'，'.join(area_desc_parts[:10])}{'等' if len(area_desc_parts) > 10 else ''}"
        text_parts.append(area_text)

    if text_parts:
        combined_text = "。".join(text_parts) + "。"
        doc.add_paragraph(combined_text)

    # AI 分析（如果启用，直接追加段落）
    if config.use_ai:
        _add_ai_town_analysis(doc, stats, config)

    doc.add_paragraph()

    # 乡镇统计表
    doc.add_heading("4.1 乡镇统计表", level=2)
    _add_town_table(doc, stats)

    # 乡镇对比图表
    if config.include_town_chart:
        town_chart = make_town_comparison_chart(stats.town_stats, stats.attr_name)
        _insert_image(doc, town_chart, config.image_width_mm, "图：各乡镇均值对比")

    if config.include_stack_chart:
        grade_order = get_grade_order(stats.attr_key)

        stack_chart = make_town_grade_stack_chart(
            stats.town_stats,
            grade_order,
            stats.attr_name,
        )
        _insert_image(doc, stack_chart, config.image_width_mm, "图：各乡镇等级占比分布")


def _add_ai_town_analysis(
    doc: Document, stats: AttributeStats, config: ReportConfig
) -> None:
    """添加AI乡镇分析"""
    from app.core.ai import generate_town_analysis

    df = stats.town_stats

    # 构建乡镇数据摘要
    town_data = {}
    for _, row in df.iterrows():
        town = row["乡镇"]
        town_data[town] = {
            "samples": int(row["样点数"]),
            "area": row["面积"],
            "mean": row["均值"] if row["均值"] == row["均值"] else None,
        }

    try:
        analysis = generate_town_analysis(
            attr_name=stats.attr_name,
            unit=stats.unit,
            town_data=town_data,
            region_name=config.region_name,
            provider=config.ai_provider,
        )
        doc.add_paragraph(analysis)
    except Exception as e:
        doc.add_paragraph(f"[乡镇分析生成失败: {e}]")


def _add_town_table(doc: Document, stats: AttributeStats) -> None:
    """添加乡镇统计表"""
    df = stats.town_stats
    grade_order = get_grade_order(stats.attr_key)

    # 基本表头
    headers = ["乡镇", "样点数", "面积(亩)", "均值"]
    # 添加等级列
    headers.extend([f"{g}占比" for g in grade_order])

    rows = []
    for _, row in df.iterrows():
        row_data = [
            row["乡镇"],
            str(int(row["样点数"])),
            f"{row['面积']:.1f}",
            f"{row['均值']:.2f}" if row["均值"] == row["均值"] else "-",
        ]
        # 等级占比
        for g in grade_order:
            pct = row.get(f"{g}_pct", 0)
            row_data.append(f"{pct:.1f}%")
        rows.append(row_data)

    _insert_table(doc, headers, rows)


# ============================================================================
# 基于 Excel 结果文件生成报告的函数
# ============================================================================


def _insert_excel_table(doc: Document, table_data: list[list[str]]) -> None:
    """将 Excel 表格数据直接插入 Word 文档

    Args:
        doc: Word 文档对象
        table_data: 二维列表，每行是一个列表
    """
    if not table_data:
        doc.add_paragraph("无数据")
        return

    # 过滤掉完全空白的行
    table_data = [row for row in table_data if any(cell.strip() for cell in row)]
    if not table_data:
        doc.add_paragraph("无数据")
        return

    # 确定最大列数
    max_cols = max(len(row) for row in table_data)

    # 创建表格
    table = doc.add_table(rows=len(table_data), cols=max_cols)
    table.style = "Table Grid"

    # 填充数据
    for row_idx, row_data in enumerate(table_data):
        for col_idx, cell_value in enumerate(row_data):
            if col_idx < max_cols:
                cell = table.rows[row_idx].cells[col_idx]
                cell.text = cell_value
                # 居中对齐
                for para in cell.paragraphs:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()


def generate_report_from_excel(
    excel_path: Path,
    attr_name: str,
    config: ReportConfig | None = None,
    output_path: Path | None = None,
) -> bytes:
    """基于 Excel 结果文件生成单个属性的分析报告

    直接从属性图数据处理生成的 Excel 文件中读取表格，插入到 Word 报告中。

    Args:
        excel_path: Excel 结果文件路径
        attr_name: 属性名称（如"有机质"）
        config: 报告配置
        output_path: 输出路径

    Returns:
        bytes: Word 文档字节数据
    """
    from app.topics.attribute_map.excel_reader import (
        read_excel_sheet_as_table,
        read_excel_stats,
    )

    if config is None:
        config = ReportConfig()

    # 从 Excel 读取统计数据（用于生成文字描述）
    excel_stats = read_excel_stats(excel_path, attr_name)
    if excel_stats is None:
        raise ValueError(f"无法从 Excel 文件读取属性 '{attr_name}' 的数据")

    doc = Document()

    # 文档标题
    title = f"{config.region_name}{excel_stats.attr_name}分析报告"
    title_para = doc.add_heading(title, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 一、总体情况
    doc.add_heading("一、总体情况", level=1)

    # 生成总体情况描述文字
    _add_overall_description(doc, excel_stats, config)

    # 溯源分析（AI生成）
    if config.use_ai:
        _add_ai_traceability_analysis_from_excel(doc, excel_stats, config)

    # 直接插入总体情况表格
    doc.add_heading("1.1 总体情况统计表", level=2)
    overall_table = read_excel_sheet_as_table(excel_path, f"{attr_name}总体情况")
    _insert_excel_table(doc, overall_table)

    # 二、土地利用类型分析
    doc.add_heading("二、土地利用类型分析", level=1)

    # 生成土地利用描述文字
    _add_land_use_description(doc, excel_stats, config)

    if config.use_ai:
        _add_ai_land_use_analysis_from_excel(doc, excel_stats, config)

    # 直接插入土地利用类型表格
    doc.add_heading("2.1 土地利用类型统计表", level=2)
    land_use_table = read_excel_sheet_as_table(
        excel_path, f"{attr_name}不同土地利用类型"
    )
    _insert_excel_table(doc, land_use_table)

    # 三、土壤类型分析
    doc.add_heading("三、土壤类型分析", level=1)

    # 生成土壤类型描述文字
    _add_soil_type_description(doc, excel_stats, config)

    if config.use_ai:
        _add_ai_soil_type_analysis_from_excel(doc, excel_stats, config)

    # 直接插入土壤类型表格
    doc.add_heading("3.1 土壤类型统计表", level=2)
    soil_type_table = read_excel_sheet_as_table(excel_path, f"{attr_name}分土壤类型")
    _insert_excel_table(doc, soil_type_table)

    # 四、乡镇分析
    doc.add_heading("四、乡镇分析", level=1)

    # 生成乡镇描述文字
    _add_town_description(doc, excel_stats, config)

    if config.use_ai:
        _add_ai_town_analysis_from_excel(doc, excel_stats, config)

    # 直接插入乡镇统计表格
    doc.add_heading("4.1 乡镇统计表", level=2)
    town_table = read_excel_sheet_as_table(excel_path, f"{attr_name}乡镇统计")
    _insert_excel_table(doc, town_table)

    # 生成时间
    doc.add_paragraph()
    time_para = doc.add_paragraph()
    time_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = time_para.add_run(f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    run.font.size = Pt(9)
    run.font.italic = True

    # 保存
    return _save_document(doc, output_path)


def _add_overall_description(doc: Document, stats, config: ReportConfig) -> None:
    """添加总体情况描述文字"""
    grade_config = SOIL_ATTR_CONFIG.get(stats.attr_key, {})
    levels = grade_config.get("levels", [])

    # 样点分析文本
    sample_parts = [
        f"{config.region_name}{stats.attr_name}样点统计分析：共采集有效样点{stats.sample_total}个，"
        f"样点{stats.attr_name}平均值为{stats.sample_mean:.2f}{stats.unit}，"
        f"中位值为{stats.sample_median:.2f}{stats.unit}，"
        f"数值范围为{stats.sample_min:.2f}~{stats.sample_max:.2f}{stats.unit}"
    ]

    if levels and stats.sample_total > 0:
        grade_desc_parts = []
        for _threshold, grade_name, _desc in levels:
            count = stats.grade_sample_counts.get(grade_name, 0)
            if count > 0:
                pct = count / stats.sample_total * 100
                grade_desc_parts.append(f"{grade_name}{count}个（占{pct:.1f}%）")
        if grade_desc_parts:
            sample_parts.append(
                f"，按等级划分，样点分布为：{'，'.join(grade_desc_parts)}"
            )

    sample_parts.append("。")
    doc.add_paragraph("".join(sample_parts))

    # 制图分析文本
    if stats.area_total > 0:
        area_parts = [
            f"{config.region_name}{stats.attr_name}制图统计分析：制图总面积为{stats.area_total:.2f}亩，"
            f"加权平均值为{stats.area_mean:.2f}{stats.unit}，"
            f"中位值为{stats.area_median:.2f}{stats.unit}，"
            f"数值范围为{stats.area_min:.2f}~{stats.area_max:.2f}{stats.unit}"
        ]

        total_area = sum(stats.grade_area_sums.values())
        if total_area > 0:
            area_grade_parts = []
            for _threshold, grade_name, _desc in levels:
                area_sum = stats.grade_area_sums.get(grade_name, 0)
                if area_sum > 0:
                    pct = area_sum / total_area * 100
                    area_grade_parts.append(
                        f"{grade_name}{area_sum:.1f}亩（占{pct:.1f}%）"
                    )
            if area_grade_parts:
                area_parts.append(
                    f"，按等级划分，面积分布为：{'，'.join(area_grade_parts)}"
                )

        area_parts.append("。")
        doc.add_paragraph("".join(area_parts))


def _add_land_use_description(doc: Document, stats, config: ReportConfig) -> None:
    """添加土地利用类型描述文字"""
    df = stats.land_use_stats

    if df.empty:
        doc.add_paragraph("数据中未包含土地利用类型信息。")
        return

    text_parts = []

    if "一级" in df.columns:
        df_filtered = df[df["二级"] != "合计"]
        if not df_filtered.empty:
            # 样点统计
            sample_summary = df_filtered.groupby("一级").agg(
                {"样点数量": "sum", "样点均值": "mean"}
            )
            sample_desc_parts = []
            for land_type, row in sample_summary.iterrows():
                if row["样点数量"] > 0:
                    sample_desc_parts.append(
                        f"{land_type}{int(row['样点数量'])}个（均值{row['样点均值']:.2f}{stats.unit}）"
                    )
            if sample_desc_parts:
                text_parts.append(
                    f"按土地利用类型，样点分布为：{'，'.join(sample_desc_parts)}"
                )

            # 面积统计
            area_summary = df_filtered.groupby("一级").agg(
                {"制图面积": "sum", "制图均值": "mean"}
            )
            area_desc_parts = []
            for land_type, row in area_summary.iterrows():
                if row["制图面积"] > 0:
                    area_desc_parts.append(
                        f"{land_type}{row['制图面积']:.1f}亩（均值{row['制图均值']:.2f}{stats.unit}）"
                    )
            if area_desc_parts:
                text_parts.append(
                    f"按土地利用类型，面积分布为：{'，'.join(area_desc_parts)}"
                )

    if text_parts:
        combined_text = "。".join(text_parts) + "。"
        doc.add_paragraph(combined_text)


def _add_soil_type_description(doc: Document, stats, config: ReportConfig) -> None:
    """添加土壤类型描述文字"""
    df = stats.soil_type_stats

    if df.empty:
        doc.add_paragraph("数据中未包含土壤类型信息。")
        return

    text_parts = []
    df_filtered = df[df["TS"] != "合计"]

    if not df_filtered.empty:
        if "sample_count" in df_filtered.columns:
            sample_summary = df_filtered.groupby("YL").agg(
                {"sample_count": "sum", "sample_mean": "mean"}
            )
            sample_desc_parts = []
            for yl, row in sample_summary.iterrows():
                if row["sample_count"] > 0:
                    sample_desc_parts.append(
                        f"{yl}{int(row['sample_count'])}个（均值{row['sample_mean']:.2f}{stats.unit}）"
                    )
            if sample_desc_parts:
                text_parts.append(
                    f"按土壤亚类，样点分布为：{'，'.join(sample_desc_parts)}"
                )

        if "area_sum" in df_filtered.columns:
            area_summary = df_filtered.groupby("YL").agg(
                {"area_sum": "sum", "area_mean": "mean"}
            )
            area_desc_parts = []
            for yl, row in area_summary.iterrows():
                if row["area_sum"] > 0:
                    area_desc_parts.append(
                        f"{yl}{row['area_sum']:.1f}亩（均值{row['area_mean']:.2f}{stats.unit}）"
                    )
            if area_desc_parts:
                text_parts.append(
                    f"按土壤亚类，面积分布为：{'，'.join(area_desc_parts)}"
                )

    if text_parts:
        combined_text = "。".join(text_parts) + "。"
        doc.add_paragraph(combined_text)


def _add_town_description(doc: Document, stats, config: ReportConfig) -> None:
    """添加乡镇描述文字"""
    df = stats.town_stats

    if df.empty:
        doc.add_paragraph("数据中未包含乡镇信息。")
        return

    text_parts = []

    # 样点描述
    sample_desc_parts = []
    for _, row in df.iterrows():
        town = row["乡镇"]
        count = int(row.get("样点数", 0))
        mean_val = row.get("均值")
        if count > 0:
            mean_str = (
                f"均值{mean_val:.2f}{stats.unit}"
                if mean_val and mean_val == mean_val
                else ""
            )
            sample_desc_parts.append(
                f"{town}{count}个{'（' + mean_str + '）' if mean_str else ''}"
            )

    if sample_desc_parts:
        sample_text = f"按乡镇，样点分布为：{'，'.join(sample_desc_parts[:10])}{'等' if len(sample_desc_parts) > 10 else ''}"
        text_parts.append(sample_text)

    # 面积描述
    area_desc_parts = []
    for _, row in df.iterrows():
        town = row["乡镇"]
        area = row.get("面积", 0)
        if area > 0:
            area_desc_parts.append(f"{town}{area:.1f}亩")

    if area_desc_parts:
        area_text = f"按乡镇，面积分布为：{'，'.join(area_desc_parts[:10])}{'等' if len(area_desc_parts) > 10 else ''}"
        text_parts.append(area_text)

    if text_parts:
        combined_text = "。".join(text_parts) + "。"
        doc.add_paragraph(combined_text)


def generate_multi_report_from_excel(
    excel_path: Path,
    attr_names: list[str],
    config: ReportConfig | None = None,
    output_path: Path | None = None,
) -> bytes:
    """基于 Excel 结果文件生成多属性综合分析报告

    Args:
        excel_path: Excel 结果文件路径
        attr_names: 属性名称列表
        config: 报告配置
        output_path: 输出路径

    Returns:
        bytes: Word 文档字节数据
    """
    from app.topics.attribute_map.excel_reader import read_excel_stats

    if config is None:
        config = ReportConfig()

    # 读取所有属性的统计数据
    excel_stats_list = []
    for attr_name in attr_names:
        stats = read_excel_stats(excel_path, attr_name)
        if stats:
            excel_stats_list.append(stats)

    if not excel_stats_list:
        raise ValueError("无法从 Excel 文件读取任何属性数据")

    doc = Document()

    # 文档标题
    title = f"{config.region_name}土壤属性综合分析报告"
    title_para = doc.add_heading(title, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 报告概述
    doc.add_heading("一、报告概述", level=1)
    intro_para = doc.add_paragraph()
    intro_para.add_run(
        f"本报告基于{config.region_name}{config.survey_year}年土壤普查数据，"
        f"对{len(excel_stats_list)}个土壤属性进行统计分析。"
    )

    attr_list = "、".join([s.attr_name for s in excel_stats_list])
    doc.add_paragraph(f"分析属性包括：{attr_list}。")

    # 各属性分析
    for i, excel_stats in enumerate(excel_stats_list, 1):
        doc.add_page_break()
        _add_attribute_section_from_excel_simple(
            doc, excel_path, excel_stats, config, section_num=i + 1
        )

    # 生成时间
    doc.add_paragraph()
    time_para = doc.add_paragraph()
    time_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = time_para.add_run(f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    run.font.size = Pt(9)
    run.font.italic = True

    return _save_document(doc, output_path)


def _add_overall_situation_from_excel(
    doc: Document, stats, config: ReportConfig
) -> None:
    """添加总体情况章节（从Excel数据）"""

    doc.add_heading("一、总体情况", level=1)

    grade_config = SOIL_ATTR_CONFIG.get(stats.attr_key, {})
    levels = grade_config.get("levels", [])

    # 构建样点分析文本
    sample_parts = [
        f"{config.region_name}{stats.attr_name}样点统计分析：共采集有效样点{stats.sample_total}个，"
        f"样点{stats.attr_name}平均值为{stats.sample_mean:.2f}{stats.unit}，"
        f"中位值为{stats.sample_median:.2f}{stats.unit}，"
        f"数值范围为{stats.sample_min:.2f}~{stats.sample_max:.2f}{stats.unit}"
    ]

    # 等级分布描述
    if levels and stats.sample_total > 0:
        grade_desc_parts = []
        for _threshold, grade_name, _desc in levels:
            count = stats.grade_sample_counts.get(grade_name, 0)
            if count > 0:
                pct = count / stats.sample_total * 100
                grade_desc_parts.append(f"{grade_name}{count}个（占{pct:.1f}%）")
        if grade_desc_parts:
            sample_parts.append(
                f"，按等级划分，样点分布为：{'，'.join(grade_desc_parts)}"
            )

    sample_parts.append("。")
    doc.add_paragraph("".join(sample_parts))

    # 构建制图分析文本
    if stats.area_total > 0:
        area_parts = [
            f"{config.region_name}{stats.attr_name}制图统计分析：制图总面积为{stats.area_total:.2f}亩，"
            f"加权平均值为{stats.area_mean:.2f}{stats.unit}，"
            f"中位值为{stats.area_median:.2f}{stats.unit}，"
            f"数值范围为{stats.area_min:.2f}~{stats.area_max:.2f}{stats.unit}"
        ]

        total_area = sum(stats.grade_area_sums.values())
        if total_area > 0:
            area_grade_parts = []
            for _threshold, grade_name, _desc in levels:
                area_sum = stats.grade_area_sums.get(grade_name, 0)
                if area_sum > 0:
                    pct = area_sum / total_area * 100
                    area_grade_parts.append(
                        f"{grade_name}{area_sum:.1f}亩（占{pct:.1f}%）"
                    )
            if area_grade_parts:
                area_parts.append(
                    f"，按等级划分，面积分布为：{'，'.join(area_grade_parts)}"
                )

        area_parts.append("。")
        doc.add_paragraph("".join(area_parts))
    else:
        doc.add_paragraph("暂无制图数据。")

    # 溯源分析（AI生成）
    if config.use_ai:
        _add_ai_traceability_analysis_from_excel(doc, stats, config)

    doc.add_paragraph()

    # 总体情况统计表
    doc.add_heading("1.1 总体情况统计表", level=2)
    _add_overall_table_from_excel(doc, stats)

    # 等级分布图表
    area_total = sum(stats.grade_area_sums.values())
    if config.include_pie_chart and area_total > 0:
        pie_data = make_grade_pie_chart(
            stats.grade_area_sums,
            f"{stats.attr_name}等级分布",
        )
        _insert_image(
            doc, pie_data, config.image_width_mm, f"图：{stats.attr_name}等级面积分布"
        )

    if config.include_bar_chart and area_total > 0:
        bar_data = make_grade_bar_chart(
            stats.grade_area_sums,
            f"{stats.attr_name}分级面积",
            ylabel="面积(亩)",
        )
        _insert_image(
            doc, bar_data, config.image_width_mm, f"图：{stats.attr_name}各等级面积对比"
        )


def _add_ai_traceability_analysis_from_excel(
    doc: Document, stats, config: ReportConfig
) -> None:
    """添加AI溯源分析（从Excel数据）"""
    from app.core.ai import generate_traceability_analysis

    total = sum(stats.grade_area_sums.values())
    grade_pct = {}
    if total > 0:
        grade_pct = {k: v / total * 100 for k, v in stats.grade_area_sums.items()}

    try:
        analysis = generate_traceability_analysis(
            attr_name=stats.attr_name,
            unit=stats.unit,
            sample_mean=stats.sample_mean,
            area_mean=stats.area_mean,
            grade_distribution=grade_pct,
            region_name=config.region_name,
            provider=config.ai_provider,
        )
        doc.add_paragraph(analysis)
    except Exception as e:
        doc.add_paragraph(f"[溯源分析生成失败: {e}]")


def _add_overall_table_from_excel(doc: Document, stats) -> None:
    """添加总体情况统计表（从Excel数据）"""
    grade_config = SOIL_ATTR_CONFIG.get(stats.attr_key, {})
    levels = grade_config.get("levels", [])

    headers = ["等级", "值域", "样点数", "样点占比", "面积(亩)", "面积占比"]
    rows = []

    sample_total = stats.sample_total
    area_total = sum(stats.grade_area_sums.values())

    for _threshold, grade_name, desc in levels:
        sample_count = stats.grade_sample_counts.get(grade_name, 0)
        area_sum = stats.grade_area_sums.get(grade_name, 0)

        sample_pct = sample_count / sample_total * 100 if sample_total > 0 else 0
        area_pct = area_sum / area_total * 100 if area_total > 0 else 0

        rows.append(
            [
                grade_name,
                desc,
                str(sample_count),
                f"{sample_pct:.1f}%",
                f"{area_sum:.1f}",
                f"{area_pct:.1f}%",
            ]
        )

    _insert_table(doc, headers, rows)

    summary = doc.add_paragraph()
    summary.add_run(
        f"合计：样点{sample_total}个，面积{area_total:.1f}亩；"
        f"均值：样点{stats.sample_mean:.2f}{stats.unit}，制图{stats.area_mean:.2f}{stats.unit}"
    )


def _add_land_use_analysis_from_excel(
    doc: Document, stats, config: ReportConfig
) -> None:
    """添加土地利用类型分析章节（从Excel数据）"""
    doc.add_heading("二、土地利用类型分析", level=1)

    df = stats.land_use_stats

    if df.empty:
        doc.add_paragraph("数据中未包含土地利用类型信息，无法进行分类分析。")
        return

    # 构建分析文本
    text_parts = []

    # 按一级分类汇总
    if "一级" in df.columns:
        # 过滤掉"合计"行
        df_filtered = df[df["二级"] != "合计"]

        if not df_filtered.empty:
            # 样点统计
            sample_summary = df_filtered.groupby("一级").agg(
                {"样点数量": "sum", "样点均值": "mean"}
            )
            sample_desc_parts = []
            for land_type, row in sample_summary.iterrows():
                if row["样点数量"] > 0:
                    sample_desc_parts.append(
                        f"{land_type}{int(row['样点数量'])}个（均值{row['样点均值']:.2f}{stats.unit}）"
                    )
            if sample_desc_parts:
                text_parts.append(
                    f"按土地利用类型，样点分布为：{'，'.join(sample_desc_parts)}"
                )

            # 面积统计
            area_summary = df_filtered.groupby("一级").agg(
                {"制图面积": "sum", "制图均值": "mean"}
            )
            area_desc_parts = []
            for land_type, row in area_summary.iterrows():
                if row["制图面积"] > 0:
                    area_desc_parts.append(
                        f"{land_type}{row['制图面积']:.1f}亩（均值{row['制图均值']:.2f}{stats.unit}）"
                    )
            if area_desc_parts:
                text_parts.append(
                    f"按土地利用类型，面积分布为：{'，'.join(area_desc_parts)}"
                )

    if text_parts:
        combined_text = "。".join(text_parts) + "。"
        doc.add_paragraph(combined_text)
    else:
        doc.add_paragraph("数据中无土地利用类型统计信息。")

    # AI 分析
    if config.use_ai:
        _add_ai_land_use_analysis_from_excel(doc, stats, config)

    doc.add_paragraph()

    # 土地利用类型统计表
    doc.add_heading("2.1 土地利用类型统计表", level=2)
    _add_land_use_table_from_excel(doc, stats)


def _add_ai_land_use_analysis_from_excel(
    doc: Document, stats, config: ReportConfig
) -> None:
    """添加AI土地利用类型分析（从Excel数据）"""
    from app.core.ai import generate_land_use_analysis

    df = stats.land_use_stats
    land_use_data = {}

    if not df.empty and "一级" in df.columns:
        df_filtered = df[df["二级"] != "合计"]
        grouped = df_filtered.groupby("一级").agg(
            {"制图面积": "sum", "制图均值": "mean"}
        )
        for land_type, row in grouped.iterrows():
            land_use_data[land_type] = {
                "area": row["制图面积"],
                "mean": row["制图均值"],
            }

    try:
        analysis = generate_land_use_analysis(
            attr_name=stats.attr_name,
            unit=stats.unit,
            land_use_data=land_use_data,
            region_name=config.region_name,
            provider=config.ai_provider,
        )
        doc.add_paragraph(analysis)
    except Exception as e:
        doc.add_paragraph(f"[土地利用分析生成失败: {e}]")


def _add_land_use_table_from_excel(doc: Document, stats) -> None:
    """添加土地利用类型统计表（从Excel数据）"""
    df = stats.land_use_stats

    headers = ["土地利用类型", "样点数", "样点均值", "面积(亩)", "面积均值"]
    rows = []

    if not df.empty and "一级" in df.columns:
        # 过滤掉"合计"行
        df_filtered = df[df["二级"] != "合计"]
        grouped = df_filtered.groupby("一级").agg(
            {
                "样点数量": "sum",
                "样点均值": "mean",
                "制图面积": "sum",
                "制图均值": "mean",
            }
        )

        for land_type, row in grouped.iterrows():
            rows.append(
                [
                    land_type,
                    str(int(row["样点数量"])),
                    f"{row['样点均值']:.2f}" if row["样点均值"] else "-",
                    f"{row['制图面积']:.1f}",
                    f"{row['制图均值']:.2f}" if row["制图均值"] else "-",
                ]
            )

    _insert_table(doc, headers, rows)


def _add_soil_type_analysis_from_excel(
    doc: Document, stats, config: ReportConfig
) -> None:
    """添加土壤类型分析章节（从Excel数据）"""
    doc.add_heading("三、土壤类型分析", level=1)

    df = stats.soil_type_stats

    if df.empty:
        doc.add_paragraph("数据中未包含土壤类型信息（YL/TS），无法进行分类分析。")
        return

    # 构建分析文本
    text_parts = []

    # 过滤掉"合计"行
    df_filtered = df[df["TS"] != "合计"]

    if not df_filtered.empty:
        # 样点统计
        if "sample_count" in df_filtered.columns:
            sample_summary = df_filtered.groupby("YL").agg(
                {"sample_count": "sum", "sample_mean": "mean"}
            )
            sample_desc_parts = []
            for yl, row in sample_summary.iterrows():
                if row["sample_count"] > 0:
                    sample_desc_parts.append(
                        f"{yl}{int(row['sample_count'])}个（均值{row['sample_mean']:.2f}{stats.unit}）"
                    )
            if sample_desc_parts:
                text_parts.append(
                    f"按土壤亚类，样点分布为：{'，'.join(sample_desc_parts)}"
                )

        # 面积统计
        if "area_sum" in df_filtered.columns:
            area_summary = df_filtered.groupby("YL").agg(
                {"area_sum": "sum", "area_mean": "mean"}
            )
            area_desc_parts = []
            for yl, row in area_summary.iterrows():
                if row["area_sum"] > 0:
                    area_desc_parts.append(
                        f"{yl}{row['area_sum']:.1f}亩（均值{row['area_mean']:.2f}{stats.unit}）"
                    )
            if area_desc_parts:
                text_parts.append(
                    f"按土壤亚类，面积分布为：{'，'.join(area_desc_parts)}"
                )

    if text_parts:
        combined_text = "。".join(text_parts) + "。"
        doc.add_paragraph(combined_text)
    else:
        doc.add_paragraph("数据中无土壤类型统计信息。")

    # AI 分析
    if config.use_ai:
        _add_ai_soil_type_analysis_from_excel(doc, stats, config)

    doc.add_paragraph()

    # 土壤类型统计表
    doc.add_heading("3.1 土壤类型统计表", level=2)
    _add_soil_type_table_from_excel(doc, stats)


def _add_ai_soil_type_analysis_from_excel(
    doc: Document, stats, config: ReportConfig
) -> None:
    """添加AI土壤类型分析（从Excel数据）"""
    from app.core.ai import generate_soil_type_analysis

    df = stats.soil_type_stats
    soil_data = {}

    if not df.empty:
        df_filtered = df[df["TS"] != "合计"]
        if "area_sum" in df_filtered.columns:
            grouped = df_filtered.groupby("YL").agg(
                {"area_sum": "sum", "area_mean": "mean"}
            )
            for yl, row in grouped.iterrows():
                soil_data[yl] = {"area": row["area_sum"], "mean": row["area_mean"]}

    try:
        analysis = generate_soil_type_analysis(
            attr_name=stats.attr_name,
            unit=stats.unit,
            soil_data=soil_data,
            region_name=config.region_name,
            provider=config.ai_provider,
        )
        doc.add_paragraph(analysis)
    except Exception as e:
        doc.add_paragraph(f"[土壤类型分析生成失败: {e}]")


def _add_soil_type_table_from_excel(doc: Document, stats) -> None:
    """添加土壤类型统计表（从Excel数据）"""
    df = stats.soil_type_stats

    if df.empty:
        doc.add_paragraph("无土壤类型统计数据。")
        return

    headers = ["亚类", "土属", "样点数", "样点均值", "面积(亩)", "面积均值"]
    rows = []

    for _, row in df.iterrows():
        yl = row.get("YL", "-")
        ts = row.get("TS", "-")
        sample_count = int(row.get("sample_count", 0)) if row.get("sample_count") else 0
        sample_mean = f"{row['sample_mean']:.2f}" if row.get("sample_mean") else "-"
        area_sum = f"{row['area_sum']:.1f}" if row.get("area_sum") else "-"
        area_mean = f"{row['area_mean']:.2f}" if row.get("area_mean") else "-"

        rows.append([yl, ts, str(sample_count), sample_mean, area_sum, area_mean])

    _insert_table(doc, headers, rows)


def _add_town_analysis_from_excel(doc: Document, stats, config: ReportConfig) -> None:
    """添加乡镇分析章节（从Excel数据）"""
    doc.add_heading("四、乡镇分析", level=1)

    df = stats.town_stats

    if df.empty:
        doc.add_paragraph("数据中未包含乡镇信息（行政区名称），无法进行分类分析。")
        return

    # 构建分析文本
    text_parts = []

    # 样点描述
    sample_desc_parts = []
    for _, row in df.iterrows():
        town = row["乡镇"]
        count = int(row.get("样点数", 0))
        mean_val = row.get("均值")
        if count > 0:
            mean_str = (
                f"均值{mean_val:.2f}{stats.unit}"
                if mean_val and mean_val == mean_val
                else ""
            )
            sample_desc_parts.append(
                f"{town}{count}个{'（' + mean_str + '）' if mean_str else ''}"
            )

    if sample_desc_parts:
        sample_text = f"按乡镇，样点分布为：{'，'.join(sample_desc_parts[:10])}{'等' if len(sample_desc_parts) > 10 else ''}"
        text_parts.append(sample_text)

    # 面积描述
    area_desc_parts = []
    for _, row in df.iterrows():
        town = row["乡镇"]
        area = row.get("面积", 0)
        if area > 0:
            area_desc_parts.append(f"{town}{area:.1f}亩")

    if area_desc_parts:
        area_text = f"按乡镇，面积分布为：{'，'.join(area_desc_parts[:10])}{'等' if len(area_desc_parts) > 10 else ''}"
        text_parts.append(area_text)

    if text_parts:
        combined_text = "。".join(text_parts) + "。"
        doc.add_paragraph(combined_text)

    # AI 分析
    if config.use_ai:
        _add_ai_town_analysis_from_excel(doc, stats, config)

    doc.add_paragraph()

    # 乡镇统计表
    doc.add_heading("4.1 乡镇统计表", level=2)
    _add_town_table_from_excel(doc, stats)

    # 乡镇对比图表
    if config.include_town_chart and not df.empty:
        town_chart = make_town_comparison_chart(df, stats.attr_name)
        _insert_image(doc, town_chart, config.image_width_mm, "图：各乡镇均值对比")

    if config.include_stack_chart:
        grade_order = get_grade_order(stats.attr_key)

        stack_chart = make_town_grade_stack_chart(
            df,
            grade_order,
            stats.attr_name,
        )
        _insert_image(doc, stack_chart, config.image_width_mm, "图：各乡镇等级占比分布")


def _add_ai_town_analysis_from_excel(
    doc: Document, stats, config: ReportConfig
) -> None:
    """添加AI乡镇分析（从Excel数据）"""
    from app.core.ai import generate_town_analysis

    df = stats.town_stats
    town_data = {}

    for _, row in df.iterrows():
        town = row["乡镇"]
        town_data[town] = {
            "samples": int(row.get("样点数", 0)),
            "area": row.get("面积", 0),
            "mean": row.get("均值") if row.get("均值") == row.get("均值") else None,
        }

    try:
        analysis = generate_town_analysis(
            attr_name=stats.attr_name,
            unit=stats.unit,
            town_data=town_data,
            region_name=config.region_name,
            provider=config.ai_provider,
        )
        doc.add_paragraph(analysis)
    except Exception as e:
        doc.add_paragraph(f"[乡镇分析生成失败: {e}]")


def _add_town_table_from_excel(doc: Document, stats) -> None:
    """添加乡镇统计表（从Excel数据）"""
    df = stats.town_stats
    grade_order = get_grade_order(stats.attr_key)

    # 基本表头
    headers = ["乡镇", "样点数", "面积(亩)", "均值"]
    # 添加等级列
    headers.extend([f"{g}占比" for g in grade_order])

    rows = []
    for _, row in df.iterrows():
        row_data = [
            row["乡镇"],
            str(int(row.get("样点数", 0))),
            f"{row.get('面积', 0):.1f}",
            f"{row['均值']:.2f}"
            if row.get("均值") and row["均值"] == row["均值"]
            else "-",
        ]
        # 等级占比
        for g in grade_order:
            pct = row.get(f"{g}_pct", 0)
            row_data.append(f"{pct:.1f}%")
        rows.append(row_data)

    _insert_table(doc, headers, rows)


def _add_attribute_section_from_excel(
    doc: Document,
    stats,
    config: ReportConfig,
    section_num: int,
) -> None:
    """添加单个属性分析章节（从Excel数据）"""
    # 章节标题
    doc.add_heading(f"{_to_chinese_num(section_num)}、{stats.attr_name}分析", level=1)

    # 统计概要
    doc.add_heading(f"{section_num}.1 统计概要", level=2)

    para = doc.add_paragraph()
    para.add_run(
        f"{stats.attr_name}（{stats.unit}）共有{stats.sample_total}个样点，"
        f"均值为{stats.sample_mean:.2f}，中位值为{stats.sample_median:.2f}，"
        f"范围为{stats.sample_min:.2f}~{stats.sample_max:.2f}。"
    )

    if stats.area_total > 0:
        para.add_run(f"制图总面积为{stats.area_total:.2f}亩。")

    # AI 分析
    if config.use_ai:
        from app.core.ai import generate_analysis

        total = sum(stats.grade_area_sums.values())
        grade_pct = (
            {k: v / total * 100 for k, v in stats.grade_area_sums.items()}
            if total > 0
            else {}
        )

        try:
            analysis = generate_analysis(
                attr_name=stats.attr_name,
                unit=stats.unit,
                sample_total=stats.sample_total,
                sample_mean=stats.sample_mean,
                sample_median=stats.sample_median,
                sample_min=stats.sample_min,
                sample_max=stats.sample_max,
                grade_distribution=grade_pct,
                region_name=config.region_name,
                provider=config.ai_provider,
            )
            doc.add_paragraph(analysis)
        except Exception:
            pass

    # 等级分布图
    area_total = sum(stats.grade_area_sums.values())
    if config.include_pie_chart and area_total > 0:
        doc.add_heading(f"{section_num}.2 等级分布", level=2)
        pie_data = make_grade_pie_chart(
            stats.grade_area_sums,
            f"{stats.attr_name}等级分布",
        )
        _insert_image(doc, pie_data, config.image_width_mm)

    # 乡镇对比
    if config.include_town_chart and not stats.town_stats.empty:
        doc.add_heading(f"{section_num}.3 乡镇对比", level=2)
        town_chart = make_town_comparison_chart(stats.town_stats, stats.attr_name)
        _insert_image(doc, town_chart, config.image_width_mm)


def _add_attribute_section_from_excel_simple(
    doc: Document,
    excel_path: Path,
    stats,
    config: ReportConfig,
    section_num: int,
) -> None:
    """添加单个属性分析章节（从Excel数据，只包含表格无图表）

    Args:
        doc: Word文档对象
        excel_path: Excel文件路径
        stats: Excel统计数据对象
        config: 报告配置
        section_num: 章节编号
    """
    from app.topics.attribute_map.excel_reader import read_excel_sheet_as_table

    attr_name = stats.attr_name

    # 章节标题
    doc.add_heading(f"{_to_chinese_num(section_num)}、{attr_name}分析", level=1)

    # 1. 总体情况
    doc.add_heading(f"{section_num}.1 总体情况", level=2)

    # 生成总体情况描述文字
    _add_overall_description(doc, stats, config)

    # 溯源分析（AI生成）
    if config.use_ai:
        _add_ai_traceability_analysis_from_excel(doc, stats, config)

    # 直接插入总体情况表格
    doc.add_heading(f"{section_num}.1.1 总体情况统计表", level=3)
    overall_table = read_excel_sheet_as_table(excel_path, f"{attr_name}总体情况")
    _insert_excel_table(doc, overall_table)

    # 2. 土地利用类型分析
    doc.add_heading(f"{section_num}.2 土地利用类型分析", level=2)

    # 生成土地利用描述文字
    _add_land_use_description(doc, stats, config)

    if config.use_ai:
        _add_ai_land_use_analysis_from_excel(doc, stats, config)

    # 直接插入土地利用类型表格
    doc.add_heading(f"{section_num}.2.1 土地利用类型统计表", level=3)
    land_use_table = read_excel_sheet_as_table(
        excel_path, f"{attr_name}不同土地利用类型"
    )
    _insert_excel_table(doc, land_use_table)

    # 3. 土壤类型分析
    doc.add_heading(f"{section_num}.3 土壤类型分析", level=2)

    # 生成土壤类型描述文字
    _add_soil_type_description(doc, stats, config)

    if config.use_ai:
        _add_ai_soil_type_analysis_from_excel(doc, stats, config)

    # 直接插入土壤类型表格
    doc.add_heading(f"{section_num}.3.1 土壤类型统计表", level=3)
    soil_type_table = read_excel_sheet_as_table(excel_path, f"{attr_name}分土壤类型")
    _insert_excel_table(doc, soil_type_table)

    # 4. 乡镇分析
    doc.add_heading(f"{section_num}.4 乡镇分析", level=2)

    # 生成乡镇描述文字
    _add_town_description(doc, stats, config)

    if config.use_ai:
        _add_ai_town_analysis_from_excel(doc, stats, config)

    # 直接插入乡镇统计表格
    doc.add_heading(f"{section_num}.4.1 乡镇统计表", level=3)
    town_table = read_excel_sheet_as_table(excel_path, f"{attr_name}乡镇统计")
    _insert_excel_table(doc, town_table)
