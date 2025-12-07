"""Word 图表插入模块

提供向 Word 文档插入图表的功能
"""

from io import BytesIO
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Mm, Pt


def insert_image_to_document(
    doc: Document,
    image_data: bytes | Path,
    width_mm: int = 150,
    caption: str | None = None,
    alignment: str = "center",
) -> None:
    """向文档插入图片

    Args:
        doc: Word 文档对象
        image_data: 图片字节数据或路径
        width_mm: 图片宽度（毫米）
        caption: 图片标题
        alignment: 对齐方式 (left, center, right)
    """
    # 添加图片
    if isinstance(image_data, Path):
        if not image_data.exists():
            return
        pic_para = doc.add_paragraph()
        run = pic_para.add_run()
        run.add_picture(str(image_data), width=Mm(width_mm))
    elif isinstance(image_data, bytes) and len(image_data) > 0:
        img_buf = BytesIO(image_data)
        pic_para = doc.add_paragraph()
        run = pic_para.add_run()
        run.add_picture(img_buf, width=Mm(width_mm))
    else:
        return

    # 设置对齐
    align_map = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
    }
    pic_para.alignment = align_map.get(alignment, WD_ALIGN_PARAGRAPH.CENTER)

    # 添加图片标题
    if caption:
        caption_para = doc.add_paragraph()
        caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = caption_para.add_run(caption)
        run.font.size = Pt(10)
        run.font.italic = True


def insert_table_from_data(
    doc: Document,
    headers: list[str],
    rows: list[list[str]],
    title: str | None = None,
) -> None:
    """向文档插入表格

    Args:
        doc: Word 文档对象
        headers: 表头列表
        rows: 数据行列表
        title: 表格标题
    """
    if title:
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title_para.add_run(title)
        run.font.bold = True
        run.font.size = Pt(11)

    # 创建表格
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"

    # 填充表头
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        # 表头加粗
        for paragraph in header_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True

    # 填充数据
    for row_idx, row_data in enumerate(rows):
        row_cells = table.rows[row_idx + 1].cells
        for col_idx, cell_value in enumerate(row_data):
            if col_idx < len(row_cells):
                row_cells[col_idx].text = str(cell_value)


def create_document_with_charts(
    charts: list[tuple[bytes, str]],
    title: str,
    output_path: Path | None = None,
    image_width_mm: int = 150,
) -> bytes:
    """创建包含多个图表的文档

    Args:
        charts: 图表列表，每项为 (图片字节, 标题)
        title: 文档标题
        output_path: 输出路径
        image_width_mm: 图片宽度

    Returns:
        bytes: Word 文档字节数据
    """
    doc = Document()

    # 添加文档标题
    title_para = doc.add_heading(title, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 添加图表
    for chart_data, chart_title in charts:
        if chart_data and len(chart_data) > 0:
            # 添加小节标题
            doc.add_heading(chart_title, level=1)

            # 插入图片
            insert_image_to_document(
                doc,
                chart_data,
                width_mm=image_width_mm,
                alignment="center",
            )

            # 添加空行
            doc.add_paragraph()

    # 保存
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    data = buf.read()

    if output_path is not None:
        output_path.parent.mkdir(parents=True, exist_ok=True)
        output_path.write_bytes(data)

    return data
